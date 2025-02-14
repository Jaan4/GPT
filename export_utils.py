# export_utils.py
from docx import Document
from fpdf import FPDF
import json
from datetime import datetime
import os

class DocumentExporter:
    """@staticmethod
     def export_pdf(documentation: str, filename: str = None) -> str:
       
        if filename is None:
            filename = f"documentation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        # Split text into lines and add to PDF
        for line in documentation.split('\n'):
            pdf.cell(0, 10, txt=line, ln=True)
        
        output_path = os.path.join('exports', filename)
        os.makedirs('exports', exist_ok=True)
        pdf.output(output_path)
        return output_path"""
    
    def export_pdf(documentation: str, filename: str = None) -> str:
        """Export documentation to PDF format"""
        
        # If filename is not provided, generate one using current date and time
        if filename is None:
            filename = f"documentation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        
        # Create a PDF instance
        pdf = FPDF()
        pdf.add_page()
        
        # Set font for the PDF
        pdf.set_font("Arial", size=12)
        
        # Split the documentation into lines and add to PDF
        for line in documentation.split('\n'):
            pdf.cell(200, 10, txt=line, ln=True)
        
        # Ensure the exports directory exists
        output_dir = 'exports'
        os.makedirs(output_dir, exist_ok=True)
        
        # Set output file path
        output_path = os.path.join(output_dir, filename)
        
        # Output the PDF to the specified path
        pdf.output(output_path)
        
        return output_path

# Example usage
if __name__ == "__main__":
    doc = "This is an example documentation.\nIt contains multiple lines.\nEach line will be written to the PDF."
    output_file = DocumentationExporter.export_pdf(doc)
    print(f"PDF generated: {output_file}")
    @staticmethod
    def export_docx(documentation: str, filename: str = None) -> str:
        """Export documentation to DOCX format"""
        if filename is None:
            filename = f"documentation_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        doc = Document()
        doc.add_heading('Code Documentation', 0)
        
        # Add content to document
        for paragraph in documentation.split('\n\n'):
            doc.add_paragraph(paragraph)
        
        output_path = os.path.join('exports', filename)
        os.makedirs('exports', exist_ok=True)
        doc.save(output_path)
        return output_path